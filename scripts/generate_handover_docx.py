from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path


@dataclass(frozen=True)
class HandoverInfo:
    project_public_name: str
    repo_url: str
    primary_maintainer_name: str
    primary_maintainer_github: str
    co_maintainer_name: str
    co_maintainer_github: str
    hosting: str
    license_name: str
    issue_tracker: str
    deployment_target: str
    credentials_policy: str
    include_current_dbs: bool
    tone: str


def _add_heading(doc, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _add_bullets(doc, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_numbered(doc, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build_doc(info: HandoverInfo, out_path: Path) -> None:
    from docx import Document  # python-docx

    doc = Document()

    # Title
    doc.add_heading(f"{info.project_public_name} — Maintainer Handover (Co-maintainer)", level=0)
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    doc.add_paragraph(f"Repository: {info.repo_url or 'TBD (to be created)'}")

    _add_heading(doc, "1) Purpose of this document", level=1)
    doc.add_paragraph(
        "This document enables a clean, repeatable handover so a co-maintainer can safely operate, "
        "extend, and (if needed) deploy the project as an open-source repository."
    )

    _add_heading(doc, "2) Maintainers & roles", level=1)
    _add_bullets(
        doc,
        [
            f"Primary maintainer: {info.primary_maintainer_name} (GitHub: @{info.primary_maintainer_github})",
            f"Co-maintainer: {info.co_maintainer_name} (GitHub: @{info.co_maintainer_github})",
            "Target co-maintainer permission: Admin (full control).",
        ],
    )

    _add_heading(doc, "3) What the project is", level=1)
    doc.add_paragraph(
        "A lightweight web app for analyzing monthly VIIRS nighttime-lights (radiance) data for cities or coordinates. "
        "It exposes a FastAPI backend for data retrieval/processing and a static HTML/JavaScript frontend for mapping "
        "and time-series visualization."
    )
    _add_bullets(
        doc,
        [
            "Backend: FastAPI endpoints for latest-available month, city/coordinate queries, and multi-city requests.",
            "Frontend: Leaflet map + D3 timeline, animation controls, and export to Excel/CSV/JSON.",
            "Data sources: NOAA EOG VIIRS monthly composites; geocoding via OpenStreetMap Nominatim.",
            "Methodology reference: Elvidge et al. (2021) bias correction concept is referenced throughout the docs/UI.",
        ],
    )

    _add_heading(doc, "4) Project aspirations & scope", level=1)
    doc.add_paragraph("Aspirations (north star):")
    _add_bullets(
        doc,
        [
            "Make nightlights trend analysis accessible with a simple UI for cities, coordinates, and comparisons.",
            "Support research-grade runs using real NOAA data (authenticated where required).",
            "Enable caching/preprocessing for fast repeated queries and predictable server costs.",
        ],
    )
    doc.add_paragraph("Current scope (what exists today):")
    _add_bullets(
        doc,
        [
            "Single-city and multi-city analysis via API + frontend UI.",
            "Real-data-only policy: API attempts to fetch real NOAA data per request (no synthetic fallback).",
            "SQLite storage for city metadata and time-series points.",
            "Docker Compose + nginx reverse proxy (VPS-friendly).",
        ],
    )
    doc.add_paragraph("Out of scope (not guaranteed / future work):")
    _add_bullets(
        doc,
        [
            "Pixel-accurate city boundary masking (current approach is radius/bounding-box style aggregation).",
            "High-scale multi-tenant hosting (this repo is currently shaped for single VPS / small user counts).",
            "Long-term data archival strategy beyond SQLite (e.g., Postgres + object storage).",
        ],
    )

    _add_heading(doc, "5) Data sources, data storage, and real-data-only policy", level=1)
    doc.add_paragraph("A) Actual / real data sources (authoritative):")
    _add_bullets(
        doc,
        [
            "NOAA EOG VIIRS monthly composites (remote). The code targets the monthly/v10 dataset.",
            "OpenStreetMap Nominatim geocoding (remote).",
        ],
    )
    doc.add_paragraph("B) Policy: no mock/synthetic data in the open-source repo:")
    _add_bullets(
        doc,
        [
            "Mock/synthetic data generation code paths are removed from the public repository.",
            "Any demo/testing artifacts should be schema-only (no synthetic radiance time series).",
        ],
    )
    doc.add_paragraph("C) Where data is stored in this repo / workspace (SQLite):")
    _add_bullets(
        doc,
        [
            "SQLite DBs are expected to be generated locally and should not be committed (except an optional schema-only demo DB).",
        ],
    )
    doc.add_paragraph("D) City list (inputs, not “data points”):")
    _add_bullets(
        doc,
        [
            "backend/cities_data.py defines the global city set used by bulk loaders and DB populators.",
        ],
    )

    _add_heading(doc, "6) Compute, storage, memory, and traffic model", level=1)
    doc.add_paragraph("Runtime components:")
    _add_bullets(
        doc,
        [
            "Backend API: CPU for downloading/processing, RAM for raster windows/arrays, disk for SQLite + temporary downloads/caches.",
            "Frontend: static files served by nginx (or simple HTTP server).",
            "Optional bulk loader: heavy network + disk temp usage during initial population; writes results to SQLite.",
        ],
    )
    doc.add_paragraph("A) Compute (CPU):")
    _add_bullets(
        doc,
        [
            "Normal interactive usage: light-to-moderate CPU (per request processing and JSON serialization).",
            "Bulk loading: heavier CPU (raster extraction + aggregation per tile/month).",
        ],
    )
    doc.add_paragraph("B) Storage (disk):")
    _add_bullets(
        doc,
        [
            "SQLite databases store only aggregated time-series points (city-month).",
            "Temporary downloads/extractions may be large in real-data mode (hundreds of MB to multiple GB per tile/month), but should be cleaned up after processing.",
            "Docker Compose defines a persistent cache volume `viirs_cache` mapped to `/tmp/viirs_cache` for backend caching.",
        ],
    )
    doc.add_paragraph("C) Memory (RAM):")
    _add_bullets(
        doc,
        [
            "Interactive API calls: RAM depends on raster extraction window size and the underlying GeoTIFF processing.",
            "Bulk loader: RAM scales with parallel tiles; reduce parallelism if OOM. Docs mention peak ~2–3GB with ~3 parallel tiles as a target operating point.",
        ],
    )
    doc.add_paragraph("D) Network / traffic:")
    _add_bullets(
        doc,
        [
            "Outbound traffic: NOAA downloads (potentially very large), plus OSM Nominatim queries (rate limited).",
            "Inbound traffic: browser to nginx + API JSON responses (small compared to NOAA downloads).",
            "Operational note: Most cost/latency risk is NOAA data fetching during cache misses / bulk population.",
        ],
    )

    _add_heading(doc, "7) Open-source readiness checklist (GitHub)", level=1)
    doc.add_paragraph("Before making the repo public:")
    _add_numbered(
        doc,
        [
            "Ensure no secrets are present in any tracked files (credentials, tokens, cookies). Use environment variables for secrets.",
            "Add/verify license file (MIT) and a short CONTRIBUTING guide if desired.",
            "Decide what data artifacts to commit (SQLite DBs, caches) vs what users generate locally.",
            "Confirm run instructions match reality (ports, docker vs local).",
        ],
    )

    _add_heading(doc, "8) Co-maintainer (Admin) handover steps", level=1)
    doc.add_paragraph("Steps for Divya to grant Admin access to Ayush on GitHub:")
    _add_numbered(
        doc,
        [
            "Create the GitHub repository (if not already created).",
            "Go to Repository → Settings → Collaborators and teams.",
            "Invite Ayush by GitHub username and set role to Admin.",
            "Ask Ayush to accept the invitation.",
            "Optional: enable branch protection rules (e.g., require PR reviews) once stable.",
        ],
    )

    _add_heading(doc, "9) Data distribution decision: committing SQLite DBs (downsides + recommendation)", level=1)
    doc.add_paragraph("You confirmed you want to include the current DB files in the repo. Downsides to be aware of:")
    _add_bullets(
        doc,
        [
            "Repo bloat: real-data DBs can grow to tens or hundreds of MB, slowing clones and PRs.",
            "Merge conflicts: SQLite is binary; concurrent changes don’t merge cleanly.",
            "Risk of accidental sensitive content: even if intended public, DBs can accidentally include test artifacts or user-provided data later.",
            "Licensing/attribution clarity: shipping derived datasets may require clear provenance statements.",
        ],
    )
    doc.add_paragraph("Practical compromise (recommended for open source):")
    _add_bullets(
        doc,
        [
            "Commit only a schema-only demo DB (no synthetic radiance values) and exclude all other DBs via .gitignore.",
            "Provide a script/command to generate DB locally using real-data mode for research.",
            "Optionally publish larger DB artifacts as GitHub Releases assets instead of tracking them in git.",
        ],
    )

    _add_heading(doc, "10) How to run (local + VPS baseline)", level=1)
    doc.add_paragraph("Local (developer):")
    _add_bullets(
        doc,
        [
            "Backend: run FastAPI/uvicorn entrypoint; frontend calls localhost API in dev mode.",
            "Frontend: open frontend/index.html or serve the frontend folder with a simple HTTP server.",
        ],
    )
    doc.add_paragraph("VPS baseline (as described in repo docs):")
    _add_bullets(
        doc,
        [
            "nginx serves static frontend and reverse proxies `/api/` to the backend container/service.",
            "Docker Compose defines `viirs-backend` and `viirs-frontend` (nginx) services and a `viirs_cache` volume.",
        ],
    )

    _add_heading(doc, "11) Known sharp edges (good to fix early)", level=1)
    _add_bullets(
        doc,
        [
            "Dependency management: backend/requirements.txt is currently empty; add pinned dependencies for reproducible installs.",
            "Search/autocomplete: frontend expects fields not returned by backend `/search` (API results may be filtered out).",
        ],
    )

    _add_heading(doc, "12) Appendix: file map (where to look)", level=1)
    _add_bullets(
        doc,
        [
            "backend/main.py — FastAPI entrypoint and endpoints",
            "backend/noaa_viirs_service.py — per-city per-month real-data downloader + extractor",
            "backend/noaa_auth.py — NOAA EOG OAuth session handler",
            "backend/database.py — SQLite schema + operations",
            "frontend/index.html + frontend/assets/js/app.js — main UI + API calls",
            "docker-compose.yml + nginx.conf — VPS-friendly deployment skeleton",
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph("End of document.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    # Fill values from the current handover context; regenerate later if needed.
    info = HandoverInfo(
        project_public_name="VIIRS Nightlights Analysis Tool",
        repo_url="https://github.com/dso6060/viirs-nightlights-analysis-tool",
        primary_maintainer_name="Divya Sornaraja",
        primary_maintainer_github="dso6060",
        co_maintainer_name="Ayush Patnaik",
        co_maintainer_github="ayushpatnaikgit",
        hosting="GitHub",
        license_name="MIT",
        issue_tracker="GitHub Issues (in-repo)",
        deployment_target="Single VPS (nginx + backend), optionally Docker Compose",
        credentials_policy="Environment variables (never commit secrets)",
        include_current_dbs=True,
        tone="friendly",
    )

    out_path = Path("docs/handover") / "VIIRS_Maintainer_Handover.docx"
    build_doc(info, out_path)
    print(f"Wrote: {out_path.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

